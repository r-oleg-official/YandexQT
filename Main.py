import os
import sqlite3
import sys
from docx import Document
from docx.shared import RGBColor, Mm
from PyQt5 import QtCore
from PyQt5.QtWidgets import QApplication, QMainWindow, QDialog, QTableWidgetItem, QFileDialog, QInputDialog
import exam_class
from Forms.main_window import Ui_MainWindow
from Forms.user_dialog import Ui_UserDialog
from Forms.exam_result import Ui_Dialog
from Forms.about_form import Ui_AboutDialog
from Forms.result_view import Ui_ResultViewDialog
from Forms.result_window import Ui_ResultsWindow
from Forms.description import Ui_Help
from Forms.new_task import Ui_new_task
from Forms.users import Ui_Users
from functions import *


def about():
    about_program = AboutDialog()
    about_program.exec_()


# Основной класс программы
class Main23EGE(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.questions = None
        self.model_1 = QtCore.QStringListModel(self)
        self.item_list = []
        self.result = {}
        self.view = self.task = self.help = self.window = None
        self.setupUi(self)
        self.con = sqlite3.connect('./data/EGE23.db')
        self.BtnCalculate.clicked.connect(self.calculation_prepare)
        self.spinStart.setValue(2)
        self.spinEnd.setValue(29)
        self.spinMove3.setValue(0)
        self.spinMove3.setHidden(True)
        self.help_about.triggered.connect(about)
        self.description.triggered.connect(self.help_how)
        self.results_in_base.triggered.connect(self.results_view)
        self.delete_results.triggered.connect(self.delete_all_results)
        self.new_variant.triggered.connect(self.variant)
        self.add_task.triggered.connect(self.new_task_edit)
        self.users.triggered.connect(self.users_in_db)
        self.move1.activated.connect(self.onActivated)
        self.move2.activated.connect(self.onActivated)
        self.move3.activated.connect(self.onActivated)
        self.BtnCalculate_2.clicked.connect(self.startExam)

    def results_view(self):
        list_of_results, start_date, end_date = ResultsView().run()
        if list_of_results:
            start_date = start_date[8:] + '/' + start_date[5:7] + '/' + start_date[:4]
            end_date = end_date[8:] + '/' + end_date[5:7] + '/' + end_date[:4]
            self.view = ResultsWindow()
            self.view.tableWidget.setRowCount(len(list_of_results))
            self.view.tableWidget.setColumnCount((len(list_of_results[0])))
            self.view.tableWidget.setColumnWidth(0, 233)
            for i, elem in enumerate(list_of_results):
                for j, val in enumerate(elem):
                    self.view.tableWidget.setItem(i, j, QTableWidgetItem(str(val)))
            self.view.label_2.setText(f'с {start_date} по {end_date}')
            self.view.show()

    def new_task_edit(self):  # открывает окно добавления нового задания в БД
        self.task = NewTask()
        self.task.show()

    def variant(
            self):  # Формирование случайного варианта для самостоятельной работы. Результат сохраняется в файле Word
        quantity, ok_pressed = QInputDialog.getItem(
            self, "Создание варианта", "Выберите количество заданий",
            ("5", "6", "7", "8", "9", "10"), 0, False)
        if ok_pressed:
            directory = os.path.dirname('./data/')
            base_name = os.path.splitext(os.path.basename('./data/'))[0]
            path = os.path.join(directory, '{0}Вариант.docx'.format(base_name))
            save_as, ans = QFileDialog.getSaveFileName(self, 'Save as...', path, "DOCX files (*.docx)")
            self.questions = get_questions(int(quantity))
            format_questions = []
            for number, el in enumerate(self.questions):
                format_question = el[1].split('|')
                format_question = '\n'.join(format_question)
                format_questions.append((format_question, el[0]))
            if save_as != '':
                document = Document()
                section = document.sections[0]
                # левое поле в миллиметрах
                section.left_margin = Mm(20.4)
                # правое поле в миллиметрах
                section.right_margin = Mm(10)
                # верхнее поле в миллиметрах
                section.top_margin = Mm(15)
                # нижнее поле в миллиметрах
                section.bottom_margin = Mm(15)
                document.add_heading('Вариант № ')
                for el in format_questions:
                    question = el[0].split('\n')
                    document.add_paragraph(question[0],
                                           style='List Number')
                    n = 2 if 'две' in question[0] else 3
                    for i in range(n):
                        line = document.add_paragraph('')
                        line.add_run(question[i + 1]).bold = True
                    document.add_paragraph(question[-1])
                    p = document.add_paragraph('')
                    run = p.add_run('Ответ: __________________')
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run = p.add_run(str(el[-1]))
                    run.font.color.rgb = RGBColor(255, 255, 255)
                document.save(save_as)
        else:
            return

    def users_in_db(self):
        self.users = Users()
        self.users.show()

    def help_how(self):  # отображение описания программы
        self.help = Description()
        self.help.show()

    def delete_all_results(self):  # Удаление всех результатов из БД
        reply = QtWidgets.QMessageBox.question(None, 'Удаление результатов из БД',
                                               'Вы действительно хотите удалить все сохраненные результаты из базы?',
                                               QtWidgets.QMessageBox.Yes, QtWidgets.QMessageBox.No)
        if reply == QtWidgets.QMessageBox.No:
            return
        else:
            cur = self.con.cursor()
            cur.execute('''DELETE FROM Results''')
            self.con.commit()

    def onActivated(self):
        """
        Функция отображения спинбоксов в зависимости от выбора команды.
        В случае если команда не выбрана - соответствующий спинбокс не отображается. И наоборот
        """
        moves_with_spin = {self.move1: self.spinMove1, self.move2: self.spinMove2, self.move3: self.spinMove3}
        if self.sender().currentIndex() == 0 or self.sender().currentIndex() == 3:
            moves_with_spin[self.sender()].setHidden(True)
        else:
            moves_with_spin[self.sender()].setHidden(False)

    def calculation_prepare(self):
        """
        Подготовка данных для расчета. Преобразование выбранных команд в необходимые функции для расчета
        (сложение, умножение, возведение в квадрат). Установка начального и конечного значений расчета.
        И при наличии- обязательных точек и точек, которые необходимо избегать.
        В зависимости от наличия обязательных и избегаемых точек - выбирается соответствующая форма расчета результата.
        """
        if not any(x for x in [self.move1.currentIndex(), self.move2.currentIndex(), self.move3.currentIndex()]):
            commands_empty(self)
        else:
            start = self.spinStart.value()
            end = self.spinEnd.value()
            symbol_to_function = {'+': lambda x, y: x + y,
                                  '*': lambda x, y: x * y,
                                  'sqr': lambda x, y: x ** 2,
                                  None: 0}
            if self.move1.currentIndex() == 1:
                move1 = (symbol_to_function['+'], self.spinMove1.value())
            elif self.move1.currentIndex() == 2:
                move1 = (symbol_to_function['*'], self.spinMove1.value())
            elif self.move1.currentIndex() == 3:
                move1 = (symbol_to_function['sqr'], self.spinMove1.value())
            else:
                move1 = (symbol_to_function[None], self.spinMove1.value())
            if self.move2.currentIndex() == 1:
                move2 = (symbol_to_function['+'], self.spinMove2.value())
            elif self.move2.currentIndex() == 2:
                move2 = (symbol_to_function['*'], self.spinMove2.value())
            elif self.move1.currentIndex() == 3:
                move2 = (symbol_to_function['sqr'], self.spinMove2.value())
            else:
                move2 = (symbol_to_function[None], self.spinMove2.value())
            if self.move3.currentIndex() == 1:
                move3 = (symbol_to_function['+'], self.spinMove3.value())
            elif self.move3.currentIndex() == 2:
                move3 = (symbol_to_function['*'], self.spinMove3.value())
            elif self.move3.currentIndex() == 3:
                move3 = (symbol_to_function['sqr'], self.spinMove3.value())
            else:
                move3 = (symbol_to_function[None], self.spinMove3.value())
            moves = (move1, move2, move3)
            self.result = {}
            required_points = []
            if self.lineEdit.text() != '':
                try:
                    required_points = list(x for x in map(int, self.lineEdit.text().split()))
                    required_points.sort()
                except ValueError:
                    self.lineEdit.clear()
                    self.lineEdit.setFocus()
                    trajectory_error(self)
                    return None
            required_tuple = tuple(required_points)
            blocked = ()
            if self.lineEdit_2.text() != '':
                block = self.lineEdit_2.text().split()
                for el in block:
                    try:
                        blocked += int(el),
                    except ValueError:
                        self.lineEdit_2.clear()
                        self.lineEdit_2.setFocus()
                        trajectory_error(self)
                        return None

            if not required_points:
                try:
                    for i in range(start, end + 1):
                        self.result[i] = (calculate(start, i, moves, blocked))
                except RecursionError:
                    commands_error(self)
                    return None
            else:
                intermediate_result = {}
                intermediate_coefficient = 1
                intermediate_point_start = start
                while len(required_points) != 0:
                    intermediate_point_end = required_points.pop(0)
                    try:
                        for i in range(intermediate_point_start, intermediate_point_end + 1):
                            if i not in intermediate_result.keys():
                                intermediate_result[i] = calculate(intermediate_point_start, i, moves,
                                                                   blocked) * intermediate_coefficient
                            else:
                                intermediate_coefficient = intermediate_result[i]
                                intermediate_result[i] = calculate(intermediate_point_start, i, moves,
                                                                   blocked) * intermediate_coefficient
                        intermediate_coefficient = i
                        intermediate_point_start = intermediate_point_end
                    except RecursionError:
                        QtWidgets.QMessageBox.critical(self, 'Неверное выражение!',
                                                       'Неверно заданы команды, либо выражение бессмысленно!',
                                                       QtWidgets.QMessageBox.Ok)
                        return None
                for i in range(intermediate_point_start, end + 1):
                    if i not in intermediate_result.keys():
                        intermediate_result[i] = calculate(intermediate_point_start, i, moves, blocked) * \
                                                 intermediate_coefficient
                    else:
                        intermediate_coefficient = intermediate_result[i]
                        intermediate_result[i] = calculate(intermediate_point_start, i, moves,
                                                           blocked) * intermediate_coefficient
                    intermediate_result[i] = calculate(intermediate_point_start, i, moves,
                                                       blocked) * intermediate_coefficient
                self.result = intermediate_result
            self.AnswerLabel.setText(str(self.result[end]))
            for i, j in self.result.items():
                if i in blocked:
                    value_for_listview = str(i) + ': ' + str(j) + '\tX'
                elif i in required_tuple:
                    value_for_listview = str(i) + ': ' + str(j) + '\t*'
                else:
                    value_for_listview = str(i) + ': ' + str(j)
                self.item_list.append(value_for_listview)
            self.listView.setAutoScroll(True)
            self.model_1.setStringList(self.item_list)
            self.listView.setModel(self.model_1)
            self.listView.scrollToBottom()

    def startExam(self):  # Отображение окна тестирования
        self.window = exam_class.Exam()
        self.window.show()


# Конец основного класса


# Класс диалогового окна выбора пользователя для тестирования
class MyDialog(QDialog, Ui_UserDialog):
    def __init__(self):
        super(MyDialog, self).__init__()
        self.setupUi(self)
        self.con = sqlite3.connect('data/EGE23.db')
        self.buttonBox_2.accepted.connect(self.run)
        self.buttonBox_2.rejected.connect(self.cancel_input)
        self.user_name = ''
        self.run()

    """ Функция выбора имени пользователя. В  комбобокс подтягиваются данные из базы данных - таблица Users
        Если поле ввода имени не заполнено - выбирается текущий элемент комбобокса. 
        В ином случае - после окончания тестирования будет создан новый пользователь, если
        он отсутствует в базе данных.
    """

    def run(self):
        try:
            cur = self.con.cursor()
            list_of_users = cur.execute('''SELECT * FROM Users''').fetchall()
            self.users_combo.addItems(list(map(lambda x: x[1], list_of_users)))
            if self.user_name_edit.text() != '':
                self.user_name = self.user_name_edit.text()
            else:
                self.user_name = self.users_combo.currentText()
        except BaseException as error:
            print(error)
        self.exec_()
        self.user_input()

    def user_input(self):
        self.close()
        return self.user_name

    def cancel_input(self):
        self.user_name = 'Cancel'
        self.close()


class ResultDialog(QDialog, Ui_Dialog):  # Класс описания результатов экзамена.
    def __init__(self):
        super(ResultDialog, self).__init__()
        self.setupUi(self)

    def run(self, score, exam_time):  # В зависимости от результата выводятся различные сообщения в окне диалога.
        if score == 100:
            self.label_result.setStyleSheet("color: rgb(64, 159, 43)")
            self.label_result.setText('Великолепно!!!')
        elif 80 <= score < 100:
            self.label_result.setStyleSheet("color: rgb(0, 150, 0)")
            self.label_result.setText('Отличный результат!')
        elif 40 <= score < 80:
            self.label_result.setStyleSheet("color: rgb(0, 0, 255)")
            self.label_result.setText('Неплохо!')
        else:
            self.label_result.setStyleSheet("color: rgb(200, 0, 0)")
            self.label_result.setText('Упс... Все ещё будет!')
        self.label_score.setStyleSheet("color: rgb(84, 138, 255)")
        self.label_score.setText(f'Вы набрали {score} баллов')
        self.label_timescore.setText(exam_time)
        self.exec_()


class AboutDialog(QDialog, Ui_AboutDialog):  # Создание окна "О программе"
    def __init__(self):
        super(AboutDialog, self).__init__()
        self.setupUi(self)


class ResultsView(QDialog, Ui_ResultViewDialog):  # Класс для просмотра результатов тестирования
    def __init__(self):
        super(ResultsView, self).__init__()
        self.setupUi(self)
        self.start_date.setDate(QtCore.QDate.currentDate())
        self.end_date.setDate(QtCore.QDate.currentDate())
        self.con = sqlite3.connect('data/EGE23.db')
        cur = self.con.cursor()
        list_of_users = cur.execute('''SELECT * FROM Users''').fetchall()  # доступен отбор по пользователю
        self.comboBox.addItems(list(map(lambda x: x[1], list_of_users)))
        self.comboBox.addItem('Все')  # либо по всем сразу
        self.comboBox.setCurrentIndex(self.comboBox.findText('Все'))
        cur.close()

    def run(self):
        self.exec_()
        start = self.start_date.dateTime().toString('yyyy/MM/dd')  # начальная дата отбора
        end = self.end_date.dateTime().toString('yyyy/MM/dd')  # конечная дата отбора
        user = self.comboBox.currentText()
        return self.view(user, start, end)

    def view(self, user, start, end):  # выборка результатов из БД
        cur = self.con.cursor()
        if user != 'Все':  # отбор по конкретному пользователю из двух таблиц
            list_for_view = cur.execute('''SELECT Results.Result, Results.DateOfExam, Results.ExamTime, 
                                            Users.Name FROM Results 
                                                  INNER JOIN Users ON Users.Id = Results.User_id 
                                                  WHERE DateOfExam >= ? and DateOfExam <= ? and Name = ?
                                                  ''', (start, end, user)).fetchall()
        else:  # отбор по всем пользователям
            list_for_view = cur.execute('''SELECT Results.Result, Results.DateOfExam, Results.ExamTime,
                                Users.Name FROM Results, Users WHERE Users.Id = Results.User_id 
                                                              AND DateOfExam >= ? and DateOfExam <= ? 
                                                              ''', (start, end)).fetchall()
        if len(list_for_view) > 0:  # если данные найдены, формирование данных для отображения
            new_list = []
            for ind, el in enumerate(list_for_view):
                new_el = ()
                new_el += (list_for_view[ind][3],)
                new_el += (list_for_view[ind][0],)
                new_el += (list_for_view[ind][1],)
                new_el += (list_for_view[ind][2],)
                new_list.append(new_el)
            cur.close()
            self.close()
            return new_list, start, end
        else:
            QtWidgets.QMessageBox.critical(None, 'Нет данных', 'По выбранным критериям данных в базе данных нет!',
                                           QtWidgets.QMessageBox.Ok)
            self.close()
            return list_for_view, start, end


class ResultsWindow(QMainWindow, Ui_ResultsWindow):  # класс окна отображения результатов
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.pushButton.clicked.connect(self.ok)
        self.pushButton_2.clicked.connect(self.save)

    def ok(self):
        self.close()

    def save(self):  # сохранение результатов в текстовый файл
        directory = os.path.dirname('./data/')
        base_name = os.path.splitext(os.path.basename('./data/'))[0]
        path = os.path.join(directory, '{0}results.txt'.format(base_name))
        save_as, ans = QFileDialog.getSaveFileName(self, 'Save as...', path, "TXT files (*.txt)")
        if save_as != '':
            output_file = open(save_as, 'w+')
            for i in range(self.tableWidget.rowCount()):
                for j in range(self.tableWidget.columnCount()):
                    output_file.write(self.tableWidget.item(i, j).text())
                    output_file.write('\t')
                output_file.write('\n')
            output_file.close()


class NewTask(QMainWindow, Ui_new_task):  # класс окна для добавления нового задания в БД
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.answer_edit.setAlignment(QtCore.Qt.AlignRight)
        self.save_btn.clicked.connect(self.save_task)
        self.cancel_btn.clicked.connect(self.close)
        self.con = sqlite3.connect('./data/EGE23.db')

    def save_task(self):  # сохранение задания с его форматированием
        if self.task_edit.toPlainText() != '':
            text = self.task_edit.toPlainText().split('\n')  # переносы на новую строку, заменяются на '|'

            new_text = '|'.join(text)  # для корректного отображения в форме тестирования
            cur = self.con.cursor()
            sql_query = """INSERT INTO Exam (Questions, RightAnswer) VALUES
                                                   (?, ?);"""
            data = (new_text, int(self.answer_edit.text()))
            cur.execute(sql_query, data)
            self.con.commit()
            cur.close()
            QtWidgets.QMessageBox.information(self, 'Сохранение задания', 'Задание успешно сохранено!',
                                              QtWidgets.QMessageBox.Ok)
        else:
            QtWidgets.QMessageBox.critical(self, 'Ошибка', 'Пустое задание!!!')
        self.close()

    def cancel(self):
        self.close()


# Класс отображения окна Описание.
class Description(QMainWindow, Ui_Help):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.file_help = open('./data/help.txt', encoding='utf-8').readlines()
        self.show()
        self.help_descrition()

    def help_descrition(self):
        for el in self.file_help:
            self.textBrowser.setText(self.textBrowser.toPlainText() + el)


# Класс пользователей программы
class Users(QMainWindow, Ui_Users):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.con = sqlite3.connect('./data/EGE23.db')
        self.btn_ok.clicked.connect(self.ok_pressed)
        self.del_btn.clicked.connect(self.delete_user)
        self.btn_add.clicked.connect(self.add_user)
        self.load_table()
        self.show()

    def load_table(self):  # заполнение таблицы пользователей
        cur = self.con.cursor()
        cur.execute(f'''SELECT Name FROM Users ORDER BY Name''')
        results = cur.fetchall()
        self.tableWidget.setColumnWidth(0, 371)
        self.tableWidget.setRowCount(0)
        for items in range(len(results)):
            self.tableWidget.setRowCount(self.tableWidget.rowCount() + 1)
            for j, elem in enumerate(results[items]):
                self.tableWidget.setItem(items, j, QtWidgets.QTableWidgetItem(elem))

    def ok_pressed(self):
        self.close()

    def add_user(self):  # добавление пользователя в БД
        name, ok_pressed = QInputDialog.getText(
            self, "Добавление пользователя", "Введите имя пользователя")
        if not name:
            return
        cur = self.con.cursor()
        sql_query = """INSERT INTO Users (Name) VALUES (?);"""
        data = (name,)
        cur.execute(sql_query, data)
        self.con.commit()
        QtWidgets.QMessageBox.information(self, 'Успешно', 'Пользователь успешно добавлен в базу!')
        self.load_table()
        cur.close()

    def delete_user(self):  # удаление выбранного пользователя из БД
        cur = self.con.cursor()
        names = self.tableWidget.selectedItems()
        for item in names:
            name = item.text()
            user_id = cur.execute("""SELECT Id FROM Users WHERE Name = ?""", (name,)).fetchone()
            cur.execute("""DELETE FROM Results WHERE User_id = ?""",
                        (user_id[0],)).fetchall()  # удаляются все результаты пользователя
            sql_query = """DELETE from Users where Id = ?"""  # затем сам пользователь
            cur.execute(sql_query, (user_id[0],))
        self.con.commit()
        QtWidgets.QMessageBox.information(self, 'Внимание', 'Пользователь удален!')
        self.load_table()
        cur.close()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Main23EGE()
    ex.show()
    sys.exit(app.exec_())
